"""
report.py - Monitoring Document Report Generator
Generates a professional Word report (.docx) from Supabase data
"""

import os
import sys
import json
import requests
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============================================================
# SUPABASE CONFIGURATION
# ============================================================
SUPABASE_URL = 'https://optezkkqvrordihanvcw.supabase.co'
SUPABASE_ANON_KEY = 'eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJpc3MiOiJzdXBhYmFzZSIsInJlZiI6Im9wdGV6a2txdnJvcmRpaGFudmN3Iiwicm9sZSI6ImFub24iLCJpYXQiOjE3ODY0OTE3ODYsImV4cCI6MjEwMjA2Nzc4Nn0.3PK7GuMHmfRDf63y4sGyOixk6DX3AS4KV3fGvNJiS0M'

# ============================================================
# SUPABASE HELPER FUNCTIONS
# ============================================================

def fetch_from_supabase(table, params=None):
    """
    Fetch data from Supabase table
    """
    headers = {
        'apikey': SUPABASE_ANON_KEY,
        'Authorization': f'Bearer {SUPABASE_ANON_KEY}',
        'Content-Type': 'application/json'
    }
    
    url = f"{SUPABASE_URL}/rest/v1/{table}"
    
    if params:
        query_parts = []
        for key, value in params.items():
            if isinstance(value, str):
                query_parts.append(f"{key}=eq.{value}")
            elif isinstance(value, list):
                for v in value:
                    query_parts.append(f"{key}=eq.{v}")
            else:
                query_parts.append(f"{key}=eq.{value}")
        
        if query_parts:
            url += "?" + "&".join(query_parts)
    
    try:
        response = requests.get(url, headers=headers)
        response.raise_for_status()
        return response.json()
    except requests.exceptions.RequestException as e:
        print(f"Error fetching from {table}: {e}")
        return []

def get_documents():
    """Get all documents from the documents table"""
    return fetch_from_supabase('documents', {'select': 'id,description,office_id'})

def get_monitoring_by_description(description):
    """Get monitoring records filtered by description"""
    return fetch_from_supabase('monitoring', {'description': description})

def get_office_name(office_id):
    """Get office name by ID"""
    offices = fetch_from_supabase('offices', {'id': office_id})
    if offices:
        return offices[0].get('name', 'N/A')
    return 'N/A'

# ============================================================
# WORD DOCUMENT GENERATION
# ============================================================

def set_cell_border(cell, border_color='000000', border_size=1):
    """Set border for a table cell"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    for border_name in ['top', 'left', 'bottom', 'right']:
        tag = f'w:{border_name}'
        border = OxmlElement(tag)
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(border_size * 8))
        border.set(qn('w:color'), border_color)
        tcPr.append(border)

def add_shading(cell, color='6366f1'):
    """Add background shading to a cell"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'solid')
    shading.set(qn('w:color'), color)
    shading.set(qn('w:fill'), color)
    tcPr.append(shading)

def create_report(description, office_name, records):
    """
    Create a Word document report
    """
    # Create document
    doc = Document()
    
    # ===== HEADER SECTION =====
    # Title
    title = doc.add_heading('MONITORING DOCUMENT REPORT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Document Type
    subtitle = doc.add_paragraph(f'Document Type: {description}')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.bold = True
    
    # Office
    office_para = doc.add_paragraph(f'Office: {office_name}')
    office_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    office_para.runs[0].font.size = Pt(12)
    
    # Date generated
    now = datetime.now()
    date_para = doc.add_paragraph(f'Generated: {now.strftime("%B %d, %Y at %I:%M %p")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.runs[0].font.size = Pt(11)
    date_para.runs[0].font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_paragraph()  # Spacing
    
    # ===== SUMMARY SECTION =====
    doc.add_heading('Report Summary', level=1)
    
    total_records = len(records)
    unique_dates = len(set(r.get('date') for r in records if r.get('date')))
    
    # Create summary table
    summary_table = doc.add_table(rows=1, cols=3)
    summary_table.style = 'Table Grid'
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    summary_table.columns[0].width = Inches(2.5)
    summary_table.columns[1].width = Inches(2.5)
    summary_table.columns[2].width = Inches(2.5)
    
    # Header row
    header_cells = summary_table.rows[0].cells
    header_cells[0].text = 'Total Records'
    header_cells[1].text = 'Monitoring Days'
    header_cells[2].text = 'Date Range'
    
    for cell in header_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(11)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_shading(cell, '6366f1')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_border(cell)
    
    # Data row
    data_row = summary_table.add_row().cells
    data_row[0].text = str(total_records)
    data_row[1].text = str(unique_dates)
    
    if records:
        dates = [r.get('date') for r in records if r.get('date')]
        date_range = f"{min(dates)} to {max(dates)}"
    else:
        date_range = 'N/A'
    data_row[2].text = date_range
    
    for cell in data_row:
        cell.paragraphs[0].runs[0].font.size = Pt(11)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_border(cell)
    
    doc.add_paragraph()  # Spacing
    
    # ===== DETAILED RECORDS SECTION =====
    doc.add_heading('Detailed Records', level=1)
    
    if records:
        # Create records table
        records_table = doc.add_table(rows=1, cols=3)
        records_table.style = 'Table Grid'
        records_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Set column widths
        records_table.columns[0].width = Inches(0.8)   # #
        records_table.columns[1].width = Inches(1.8)   # Date
        records_table.columns[2].width = Inches(4.4)   # Follow-up
        
        # Header row
        header_cells = records_table.rows[0].cells
        header_cells[0].text = '#'
        header_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_cells[1].text = 'Date'
        header_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_cells[2].text = 'Follow-up Respond'
        
        for cell in header_cells:
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            add_shading(cell, '6366f1')
            set_cell_border(cell)
        
        # Data rows
        for idx, record in enumerate(records, 1):
            row_cells = records_table.add_row().cells
            row_cells[0].text = str(idx)
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[1].text = record.get('date', '')
            row_cells[2].text = record.get('followup', '')
            
            for cell in row_cells:
                cell.paragraphs[0].runs[0].font.size = Pt(10)
                set_cell_border(cell)
        
        # Add total row
        total_row = records_table.add_row().cells
        total_row[0].merge(total_row[2])
        total_row[0].text = f'Total Records: {total_records}'
        total_row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        total_row[0].paragraphs[0].runs[0].font.bold = True
        total_row[0].paragraphs[0].runs[0].font.size = Pt(10)
        
        # Set background for total row
        add_shading(total_row[0], 'f1f5f9')
        set_cell_border(total_row[0])
        
    else:
        doc.add_paragraph('No records found for this description.')
        doc.add_paragraph().runs[0].font.size = Pt(11)
    
    # ===== FOOTER =====
    doc.add_paragraph()  # Spacing
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    footer_run = footer.add_run('This report is automatically generated from the Monitoring Document System.')
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(150, 150, 150)
    
    footer_run2 = footer.add_run(f'\nGenerated on {now.strftime("%B %d, %Y at %I:%M %p")}')
    footer_run2.font.size = Pt(9)
    footer_run2.font.color.rgb = RGBColor(150, 150, 150)
    
    return doc

def generate_and_save_report(description, office_name, records, filename=None):
    """
    Generate and save the report
    """
    doc = create_report(description, office_name, records)
    
    if filename is None:
        now = datetime.now()
        safe_description = description.replace(' ', '_').replace('/', '-')
        filename = f"Monitoring_Report_{safe_description}_{now.strftime('%Y-%m-%d_%H-%M')}.docx"
    
    # Ensure the reports directory exists
    reports_dir = 'reports'
    if not os.path.exists(reports_dir):
        os.makedirs(reports_dir)
    
    filepath = os.path.join(reports_dir, filename)
    doc.save(filepath)
    print(f"✅ Report saved as: {filepath}")
    return filepath

# ============================================================
# FLASK API ENDPOINT (Optional)
# ============================================================

def create_flask_app():
    """Create Flask app for API endpoint"""
    try:
        from flask import Flask, request, jsonify, send_file
        from flask_cors import CORS
        
        app = Flask(__name__)
        CORS(app)  # Enable CORS for all routes
        
        @app.route('/api/generate_report', methods=['POST', 'OPTIONS'])
        def generate_report_api():
            if request.method == 'OPTIONS':
                return jsonify({'status': 'ok'}), 200
            
            try:
                data = request.json
                description = data.get('description')
                office = data.get('office', 'N/A')
                records = data.get('records', [])
                
                print(f"📊 Generating report for: {description}")
                print(f"📝 Records: {len(records)}")
                
                if not description or not records:
                    return jsonify({
                        'success': False,
                        'message': 'Missing description or records'
                    }), 400
                
                # Generate report
                filepath = generate_and_save_report(description, office, records)
                filename = os.path.basename(filepath)
                
                return jsonify({
                    'success': True,
                    'filename': filename,
                    'download_url': f'/downloads/{filename}',
                    'message': f'Report generated successfully: {filename}'
                })
                
            except Exception as e:
                print(f"❌ Error: {e}")
                return jsonify({
                    'success': False,
                    'message': str(e)
                }), 500
        
        @app.route('/downloads/<filename>')
        def download_report(filename):
            """Download the generated report"""
            filepath = os.path.join('reports', filename)
            if os.path.exists(filepath):
                return send_file(filepath, as_attachment=True)
            return jsonify({'error': 'File not found'}), 404
        
        return app
        
    except ImportError:
        print("Flask not installed. Install with: pip install flask flask-cors")
        return None

# ============================================================
# MAIN FUNCTION
# ============================================================

def main():
    """
    Main function to generate report from command line arguments
    or start Flask server
    """
    import argparse
    
    parser = argparse.ArgumentParser(description='Generate Monitoring Report')
    parser.add_argument('--description', '-d', help='Document description')
    parser.add_argument('--server', '-s', action='store_true', help='Run as Flask server')
    parser.add_argument('--port', '-p', type=int, default=5000, help='Port for Flask server')
    
    args = parser.parse_args()
    
    if args.server:
        # Run as Flask server
        print("🚀 Starting Flask server...")
        app = create_flask_app()
        if app:
            print(f"✅ Server running on http://localhost:{args.port}")
            print("📊 API endpoint: http://localhost:{args.port}/api/generate_report")
            app.run(debug=True, port=args.port)
        else:
            print("❌ Flask not installed. Install with: pip install flask flask-cors")
        return
    
    if args.description:
        # Generate report from command line
        description = args.description
        print(f"📊 Generating report for: {description}")
        
        # Fetch data from Supabase
        print("🔍 Fetching data from Supabase...")
        
        # Get office name
        documents = get_documents()
        office_id = None
        for doc in documents:
            if doc.get('description') == description:
                office_id = doc.get('office_id')
                break
        
        if office_id:
            office_name = get_office_name(office_id)
        else:
            office_name = 'N/A'
            print(f"⚠️ No office found for description '{description}'")
        
        # Get monitoring records
        records = get_monitoring_by_description(description)
        
        if not records:
            print(f"⚠️ No monitoring records found for '{description}'")
            proceed = input("Continue with empty report? (y/n): ").strip().lower()
            if proceed != 'y':
                return
        
        print(f"✅ Found {len(records)} records")
        
        # Generate and save report
        print("📝 Generating report...")
        filepath = generate_and_save_report(description, office_name, records)
        
        print("\n" + "=" * 60)
        print("✅ Report generation complete!")
        print(f"📁 File: {filepath}")
        print("=" * 60)
    else:
        # Interactive mode
        print("=" * 60)
        print("MONITORING DOCUMENT REPORT GENERATOR")
        print("=" * 60)
        print("\nOptions:")
        print("  1. Generate report from command line")
        print("  2. Run as Flask server")
        print("\nExamples:")
        print("  python report.py --description 'Annual Report'")
        print("  python report.py --server --port 5000")
        print("\n" + "=" * 60)

if __name__ == "__main__":
    main()